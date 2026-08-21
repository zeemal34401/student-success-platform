#!/usr/bin/env python3
"""Also write a Word companion with full slide-by-slide speaker content."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "PrognosEd_Conference_Slide_Content.docx"

PRIMARY = RGBColor(0x0B, 0x6E, 0x4F)
CHARCOAL = RGBColor(0x10, 0x18, 0x28)
MUTED = RGBColor(0x47, 0x54, 0x67)

SLIDES = [
    (
        1,
        "Title Slide — PrognosEd: AI-Powered Early-Warning Student Success Platform",
        [
            "PrognosEd",
            "AI-Powered Early-Warning Student Success Platform",
            "Predictive Analytics · Intervention Intelligence · Retrieval-Augmented Advising",
            "International Conference Presentation",
        ],
        "Open with the problem of delayed intervention, then position PrognosEd as an end-to-end AI platform.",
    ),
    (
        2,
        "Presentation Outline",
        [
            "1. Problem context and research motivation",
            "2. Objectives, proposed solution, and system architecture",
            "3. AI/ML methodology: four specialized predictive models",
            "4. Key innovation — retrieval-augmented (RAG) advising chat",
            "5. Role-based decision support and live prototype demonstration",
            "6. Experimental results, impact, limitations, and future work",
            "7. Conclusions and discussion",
        ],
        "Keep this slide brief; it orients the audience for a 15–20 minute talk.",
    ),
    (
        3,
        "The Challenge: Delayed Intervention in Higher Education",
        [
            "Universities often detect academic distress after grades have already declined.",
            "Faculty workload limits continuous monitoring across large course cohorts.",
            "Attendance, LMS activity, and assessment signals remain fragmented and underused.",
            "One-size-fits-all advising fails to prioritize who needs support first.",
            "Result: reactive remediation, higher dropout risk, and weaker retention outcomes.",
            "Need: an early-warning intelligence system that predicts risk and recommends action.",
        ],
        "Emphasize why timing matters: earlier signals enable lower-cost, higher-impact support.",
    ),
    (
        4,
        "Research Objectives and Scope",
        [
            "Design an AI-enabled platform for early academic and dropout risk prediction.",
            "Deliver personalized, ranked intervention recommendations for faculty action.",
            "Provide role-based dashboards for Faculty, Admin, Department Head, and Director.",
            "Integrate multi-model ML services with graceful offline degradation.",
            "Introduce conversational RAG advising grounded in student metrics and history.",
            "Validate through a production-structured prototype with measurable model metrics.",
        ],
        "State objectives as measurable engineering and research outcomes.",
    ),
    (
        5,
        "Proposed Solution Overview",
        [
            "PrognosEd — an integrated Student Success Intelligence Platform.",
            "Ingests student engagement and academic metrics (attendance, GPA, LMS, late work).",
            "Predicts risk via four independent ML microservices.",
            "Surfaces prioritized alerts and intervention plans in role-scoped dashboards.",
            "Supports decisions with a per-student RAG advising assistant.",
            "Secured with JWT authentication, RBAC, rate limiting, and SMTP faculty onboarding.",
        ],
        "Transition from problem → solution in one clear system narrative.",
    ),
    (
        6,
        "System Architecture",
        [
            "Presentation Layer: React 19 SPA, role-based dashboards, RAG chat UI",
            "Application Layer: Express REST API, JWT sessions, SQLite persistence",
            "AI & Analytics Layer: FastAPI ML services on ports 8000–8003",
            "Knowledge Layer: TF-IDF retriever, advising playbook, chat memory",
            "Design principle: modular services, clear APIs, resilient fallbacks",
        ],
        "Show architecture diagram (docs/print-assets/01-system-architecture.png) if available.",
    ),
    (
        7,
        "Technology Stack",
        [
            "Frontend: React 19, Vite, Tailwind CSS 4, Recharts",
            "Backend: Node.js, Express 4, better-sqlite3, bcrypt, JWT",
            "ML Serving: FastAPI + uvicorn, scikit-learn / XGBoost pipelines",
            "Advising AI: Node.js TF-IDF RAG retriever with intent-aware generation",
            "Ops: Docker Compose support, GitHub Actions CI (lint, build, smoke tests)",
            "Email: Nodemailer SMTP for invitations and password reset",
        ],
        "Highlight production-structured choices, not only research notebooks.",
    ),
    (
        8,
        "AI Pipeline: From Student Metrics to Action",
        [
            "Step 1 — Collect metrics: attendance, GPA, LMS activity, late assignments",
            "Step 2 — Predict risk: academic early-warning + dropout probability",
            "Step 3 — Rank interventions: skill gaps + risk-level action templates",
            "Step 4 — Faculty action: dashboard alerts, accept/dismiss recommendations",
            "Step 5 — Continuous advising: RAG chat with per-student conversation memory",
            "Feedback loop: recommendation decisions refine operational prioritization",
        ],
        "Use workflow diagram (docs/print-assets/02-ai-workflow.png).",
    ),
    (
        9,
        "Model 1 — Academic Early-Warning Risk Prediction",
        [
            "Purpose: predict pass/fail risk early enough for meaningful intervention",
            "Signals: LMS engagement (clicks, active days) and assessment performance",
            "Evaluation design: 4-week early-warning window (not full-course leakage)",
            "Key result: ROC-AUC ≈ 0.83 on OULAD early-warning setting",
            "Explainability: SHAP confirms behavioral/academic drivers over demographics",
            "Service endpoint: FastAPI on port 8000 (ML_ACADEMIC_API_URL)",
        ],
        "Stress the 4-week constraint as methodological honesty for real intervention timing.",
    ),
    (
        10,
        "Model 2 — Dropout Risk Detection",
        [
            "Purpose: binary classification of Dropout vs Graduate outcomes",
            "Features: enrollment, curricular units, financial/academic indicators",
            "Dataset: UCI Higher Education Students Performance / Dropout",
            "Key results: 94.1% accuracy; AUC 0.973",
            "Validation: stratified 5-fold CV + held-out test reporting",
            "Service endpoint: FastAPI on port 8001 (ML_DROPOUT_API_URL)",
        ],
        "Report both accuracy and AUC; mention CV to avoid single-split optimism.",
    ),
    (
        11,
        "Model 3 — Skill-Based Intervention Recommender",
        [
            "Purpose: recommend skill-level interventions from mastery gaps",
            "Approach: rank weak skill areas from tutoring / EDM mastery tables",
            "Dataset lineage: KDD Cup 2010 Educational Data Mining challenge",
            "Key result: Precision@5 = 100% on evaluated ranking setting",
            "Integration: complements platform rule-based intervention templates",
            "Service endpoint: FastAPI on port 8002 (ML_RECOMMENDER_API_URL)",
        ],
        "Clarify that Model 3 ranks actionable skill targets, not only risk scores.",
    ),
    (
        12,
        "Model 4 — Behavioral Learning Analytics Clustering",
        [
            "Purpose: unsupervised segmentation for institutional analytics",
            "Method: k-means behavioral clustering (k = 4)",
            "Key metric: Silhouette score ≈ 0.315 (moderate, exploratory separation)",
            "Use case: cohort patterns for directors and academic administrators",
            "Supports reporting beyond individual risk lists",
            "Service endpoint: FastAPI on port 8003 (ML_ANALYTICS_API_URL)",
        ],
        "Frame clustering as exploratory analytics for leadership, not clinical diagnosis.",
    ),
    (
        13,
        "Key Innovation: Retrieval-Augmented Advising Chat",
        [
            "Per-student conversational assistant embedded in the student profile",
            "Retrieves from an advising playbook using TF-IDF (Node.js, zero heavy deps)",
            "Personalizes answers with live metrics: risk, attendance, GPA, LMS, late work",
            "Stores short conversation memory per student–advisor pair in SQLite",
            "Intent detection: outreach drafts, weekly plans, risk explanation, summaries",
            "Anti-repetition controls rotate focus and avoid recycled playbook chunks",
        ],
        "Position RAG as the bridge between model outputs and day-to-day faculty practice.",
    ),
    (
        14,
        "Role-Based Decision Support (Faculty → Leadership)",
        [
            "Faculty: course KPIs, at-risk roster, recommendations, RAG advisor",
            "Department Head: department performance and faculty-aligned views",
            "Academic Admin: institutional insights, user invites, course assignments",
            "Director / Dean: university KPIs, department comparison, risk drill-down",
            "Administrative Staff: reports and alerts for operational support",
            "Access control enforced via JWT + role-scoped service queries",
        ],
        "Show that AI insights are organizationally usable across governance levels.",
    ),
    (
        15,
        "Prototype Demonstration — Faculty Dashboard",
        [
            "Live UI demonstration of Faculty Dashboard",
            "KPI cards: enrolled students, average attendance, at-risk count, average GPA",
            "Weekly engagement trend: attendance vs LMS activity",
            "Students needing attention panel for rapid triage",
            "Screenshot asset: docs/print-assets/05-faculty-dashboard.png",
        ],
        "Narrate a real faculty morning workflow: scan KPIs → open at-risk list.",
    ),
    (
        16,
        "Prototype Demonstration — Student Risk Profile & Interventions",
        [
            "Critical-risk student profile with contributing factors",
            "ML model predictions: academic, dropout, and engagement-based risk",
            "Personalized intervention checklist for immediate action",
            "Personalized RAG advisor with conversation history",
            "Screenshot asset: docs/print-assets/06-student-profile-rag.png",
        ],
        "Walk through one student story end-to-end (e.g., high late work + low GPA).",
    ),
    (
        17,
        "Experimental Results and Model Performance",
        [
            "Model 1 (Academic early-warning): ROC-AUC ≈ 0.83 (4-week setting)",
            "Model 2 (Dropout risk): Accuracy 94.1%; AUC 0.973",
            "Model 3 (Skill recommender): Precision@5 = 100%",
            "Model 4 (Behavioral clustering): Silhouette ≈ 0.315 (k = 4)",
            "Platform resilience: heuristic fallbacks when ML services are offline",
            "Prototype coverage: dashboards, alerts, recommendations, reports, RAG chat",
        ],
        "Optionally display docs/print-assets/03-key-results.png as a visual panel.",
    ),
    (
        18,
        "Specification Compliance and Deliverables",
        [
            "✓ Faculty dashboard with at-risk prioritization",
            "✓ Risk alert system with role-scoped visibility",
            "✓ Recommendation engine (rules + ML skill diagnostics)",
            "✓ Institutional / class reporting analytics",
            "✓ Four ML microservices for prediction and analytics",
            "✓ Extended deliverables: RAG chat, SMTP onboarding, RBAC hierarchy",
        ],
        "Useful for evaluators comparing against the original project specification.",
    ),
    (
        19,
        "Security, Scalability, and Deployment",
        [
            "Security: JWT sessions, bcrypt hashing, API rate limits, CORS controls",
            "Privacy posture: role-scoped data access; demo accounts disabled in production",
            "Deployment: npm production build, Express static hosting, Docker Compose",
            "Ops readiness: health endpoints for API, ML services, and email",
            "Scale path: SQLite for pilot → PostgreSQL for institutional concurrency",
            "CI quality gates: lint, build, and smoke tests on push/PR",
        ],
        "Reassure conference audiences that the prototype is deployment-aware.",
    ),
    (
        20,
        "Limitations and Ethical Considerations",
        [
            "Some ML features are estimated from summary metrics (not raw LMS event streams).",
            "RAG uses a curated playbook; not a full institutional policy corpus or LLM.",
            "Clustering is exploratory; moderate silhouette implies interpretive caution.",
            "Predictions must support—not replace—professional academic judgment.",
            "Fairness vigilance: minimize reliance on sensitive demographic proxies.",
            "Production use requires FERPA/GDPR-aligned governance and audit logging.",
        ],
        "Honesty here builds credibility with international academic audiences.",
    ),
    (
        21,
        "Impact and Contributions",
        [
            "Shifts advising from reactive remediation to proactive early warning",
            "Unifies prediction, recommendation, and conversational guidance in one system",
            "Provides leadership visibility from course level to institutional KPIs",
            "Demonstrates modular, production-structured ML microservice integration",
            "Introduces per-student RAG memory for continuity of advising context",
            "Offers a deployable prototype pathway for university pilot programs",
        ],
        "Close the value loop: technical novelty + institutional usefulness.",
    ),
    (
        22,
        "Future Work and Research Directions",
        [
            "Live LMS/SIS ingestion (LTI, xAPI, or scheduled ETL pipelines)",
            "Embedding-based retrieval + grounded LLM synthesis with citations",
            "PostgreSQL migration, connection pooling, and horizontal scaling",
            "Automated retraining, model versioning, and monitoring (e.g., MLflow)",
            "Real-time alert digests and institutional workflow integrations",
            "Expanded fairness audits, accessibility (WCAG), and longitudinal impact studies",
        ],
        "Invite collaboration: data partnerships, pilots, and joint research.",
    ),
    (
        23,
        "Conclusion",
        [
            "PrognosEd operationalizes AI for early student-success intervention.",
            "Four ML services cover academic risk, dropout, skills, and behavioral analytics.",
            "RAG advising adds personalized, history-aware decision support for faculty.",
            "Role-based dashboards connect course action to institutional strategy.",
            "Strong prototype readiness with clear, evidence-based performance metrics.",
            "Next step: pilot deployment with live institutional data partnerships.",
        ],
        "Restate the contribution in one sentence before Q&A.",
    ),
    (
        24,
        "Acknowledgments",
        [
            "Academic supervision and research guidance for the AI/ML internship track",
            "Open datasets enabling reproducible model development",
            "OULAD · UCI Student Performance · UCI Dropout · KDD Cup 2010 EDM",
            "Open-source ecosystem: React, Express, FastAPI, scikit-learn, and related tools",
            "Institutional stakeholders who inspired practical faculty advising workflows",
        ],
        "Keep acknowledgments concise and respectful.",
    ),
    (
        25,
        "Questions & Discussion",
        [
            "Thank you for your attention",
            "PrognosEd — AI-Powered Early-Warning Student Success Platform",
            "Open floor for questions on methodology, deployment, ethics, or collaboration",
        ],
        "Prepare concise answers on datasets, fairness, RAG limits, and pilot readiness.",
    ),
]


def set_run(run, size=11, bold=False, color=CHARCOAL):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PrognosEd — International Conference Slide Content")
    set_run(r, 20, True, PRIMARY)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(
        f"Full slide-by-slide text for PowerPoint  ·  {date.today().strftime('%B %d, %Y')}"
    )
    set_run(r, 11, False, MUTED)

    n = doc.add_paragraph()
    r = n.add_run(
        "Use the bullets under each heading directly on slides. "
        "Speaker notes are guidance for oral delivery and do not need to appear on-screen."
    )
    set_run(r, 10, False, MUTED)

    for num, title, bullets, note in SLIDES:
        h = doc.add_heading(f"Slide {num}: {title}", level=1)
        for run in h.runs:
            run.font.color.rgb = PRIMARY

        for b in bullets:
            p = doc.add_paragraph(b, style="List Bullet")
            for run in p.runs:
                set_run(run, 11, False, CHARCOAL)

        p = doc.add_paragraph()
        r1 = p.add_run("Speaker note: ")
        set_run(r1, 10, True, PRIMARY)
        r2 = p.add_run(note)
        set_run(r2, 10, False, MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    print(build())
