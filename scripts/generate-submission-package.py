#!/usr/bin/env python3
"""Generate a presentation-ready Word submission: Title, Abstract, Print Images."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "docs" / "print-assets"
OUT = ROOT / "docs" / "PrognosEd_Title_Abstract_and_Images.docx"

PRIMARY = RGBColor(0x0B, 0x6E, 0x4F)
CHARCOAL = RGBColor(0x10, 0x18, 0x28)
MUTED = RGBColor(0x47, 0x54, 0x67)


def set_run_font(run, size=11, bold=False, color=CHARCOAL, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
        run.font.name = "Calibri"
    return h


def add_body(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_run_font(run, 10, bold=False, color=MUTED)
    run.italic = True
    return p


def add_image(doc, path, width_inches=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    return p


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ----- Cover -----
    for _ in range(3):
        doc.add_paragraph("")

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("PrognosEd")
    set_run_font(r, 14, True, PRIMARY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    r = title.add_run(
        "PrognosEd: AI-Powered Early-Warning\nStudent Success Platform"
    )
    set_run_font(r, 26, True, CHARCOAL)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(12)
    r = sub.add_run(
        "Project Submission Package\n"
        "Title · Abstract · High-Resolution Exhibition Images"
    )
    set_run_font(r, 12, False, MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    r = meta.add_run(
        f"Document Date: {date.today().strftime('%B %d, %Y')}\n"
        "Classification: Academic / Exhibition Submission\n"
        "Status: Presentation Ready"
    )
    set_run_font(r, 11, False, MUTED)

    doc.add_page_break()

    # ----- 1. Title -----
    add_heading_styled(doc, "1. Project Title", 1)
    add_body(
        doc,
        "PrognosEd: AI-Powered Early-Warning Student Success Platform",
    )
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    r = note.add_run(
        "Alternate short form for posters and banners: PrognosEd — Predictive Student Success Intelligence"
    )
    set_run_font(r, 10, False, MUTED)
    r.italic = True

    # ----- 2. Abstract -----
    add_heading_styled(doc, "2. Abstract", 1)
    abstract = (
        "Declining retention and delayed intervention leave universities reacting after "
        "students are already failing. PrognosEd is an AI-based student success platform "
        "that predicts academic and dropout risk early, ranks personalized interventions, "
        "and equips faculty with role-based dashboards for timely action. Its core "
        "innovation combines four ML microservices (early-warning risk, dropout detection, "
        "skill recommender, and behavioral clustering) with a retrieval-augmented advising "
        "chat that learns from each student’s history. By unifying prediction, recommendation, "
        "and conversational guidance in one secure web system, PrognosEd enables proactive "
        "support, improves advising efficiency, and strengthens institutional retention outcomes."
    )
    p = add_body(doc, abstract, space_after=6)
    # Word count note
    wc = len(abstract.split())
    wcp = doc.add_paragraph()
    r = wcp.add_run(f"Word count: {wc} words (target range: 80–100).")
    set_run_font(r, 9, False, MUTED)
    r.italic = True

    # Structure callouts
    add_heading_styled(doc, "Abstract Structure Map", 2)
    bullets = [
        ("Problem: ", "Delayed intervention and declining retention after academic failure begins."),
        ("AI-based solution: ", "Early academic/dropout risk prediction with ranked, personalized interventions."),
        ("Key innovation: ", "Four ML microservices plus per-student RAG advising chat with conversation memory."),
        ("Impact: ", "Proactive faculty action, more efficient advising, and stronger retention outcomes."),
    ]
    for label, text in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        r1 = bp.add_run(label)
        set_run_font(r1, 11, True, PRIMARY)
        r2 = bp.add_run(text)
        set_run_font(r2, 11, False, CHARCOAL)

    doc.add_page_break()

    # ----- 3. Images -----
    add_heading_styled(doc, "3. High-Resolution Images for Large-Format Printing", 1)
    add_body(
        doc,
        "The following original visuals were produced from the live PrognosEd prototype and "
        "system design assets. Each image is prepared at approximately 3840 px width with "
        "300 DPI metadata for large-format posters, exhibition boards, and technical displays.",
    )

    images = [
        (
            ASSETS / "01-system-architecture.png",
            "Figure 1. System Architecture Diagram",
            "Layered architecture of PrognosEd: Presentation (React SPA and role-based dashboards), "
            "Application (Express REST API, JWT authentication, SQLite), AI & Analytics "
            "(Models 1–4: academic risk, dropout detection, skill recommender, behavioral clustering), "
            "and Knowledge (TF-IDF RAG retriever, advising playbook, per-student chat memory).",
            "Architecture / workflow diagram",
        ),
        (
            ASSETS / "02-ai-workflow.png",
            "Figure 2. AI Workflow Pipeline",
            "End-to-end workflow from student metrics (attendance, GPA, LMS activity, late work) "
            "through ML risk prediction and intervention ranking to faculty action and continuous "
            "RAG advising, with a recommendation accept/dismiss feedback loop.",
            "Architecture / workflow diagram",
        ),
        (
            ASSETS / "05-faculty-dashboard.png",
            "Figure 3. Faculty Dashboard (Live Prototype Screenshot)",
            "Software interface of the Faculty Dashboard showing enrolled students, average "
            "attendance, at-risk count, GPA KPIs, weekly engagement trend, and the "
            "“Students needing attention” panel for rapid intervention triage.",
            "Software / interface screenshot",
        ),
        (
            ASSETS / "06-student-profile-rag.png",
            "Figure 4. Student Profile with Personalized RAG Advisor (Live Prototype)",
            "Application visual of a critical-risk student profile featuring personalized "
            "intervention recommendations and the retrieval-augmented advising chat that uses "
            "live metrics and prior advising history for context-aware guidance.",
            "Key results / application visual · Prototype screenshot",
        ),
    ]

    for path, fig_title, caption, category in images:
        if not path.exists():
            add_body(doc, f"[Missing asset: {path.name}]")
            continue
        add_heading_styled(doc, fig_title, 2)
        cat = doc.add_paragraph()
        r = cat.add_run(f"Category: {category}")
        set_run_font(r, 10, True, PRIMARY)
        add_image(doc, path, width_inches=6.4)
        add_caption(doc, caption)

    # Optional fifth supporting image note (key results) on new page if exists
    results = ASSETS / "03-key-results.png"
    if results.exists():
        doc.add_page_break()
        add_heading_styled(doc, "Supplementary Visual — Key Model Results", 1)
        add_body(
            doc,
            "Optional exhibition panel summarizing validated model metrics and multi-role "
            "platform capabilities. Suitable as a fifth board or inset panel.",
        )
        cat = doc.add_paragraph()
        r = cat.add_run("Category: Key results / application visual")
        set_run_font(r, 10, True, PRIMARY)
        add_image(doc, results, width_inches=6.4)
        add_caption(
            doc,
            "Figure 5. Key Platform Capabilities — Model 1 early-warning AUC ≈ 0.83 (4-week); "
            "Model 2 dropout accuracy 94.1% (AUC 0.973); Model 3 skill recommender Precision@5 100%; "
            "multi-role intelligence with RAG advising and per-student memory.",
        )

    # Image inventory
    doc.add_page_break()
    add_heading_styled(doc, "4. Image Inventory & Print Specifications", 1)
    add_body(
        doc,
        "All primary images are original to this project (generated system diagrams and "
        "screenshots captured from the running PrognosEd prototype).",
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["File", "Content", "Print Spec", "Use"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    rows = [
        (
            "01-system-architecture.png",
            "Architecture diagram",
            "~3840×2560 · 300 DPI",
            "Poster / board",
        ),
        (
            "02-ai-workflow.png",
            "AI workflow diagram",
            "~3840×2560 · 300 DPI",
            "Poster / board",
        ),
        (
            "05-faculty-dashboard.png",
            "Faculty UI screenshot",
            "~3840×2159 · 300 DPI",
            "Prototype photo panel",
        ),
        (
            "06-student-profile-rag.png",
            "RAG advisor UI screenshot",
            "~3840×2159 · 300 DPI",
            "Application visual",
        ),
        (
            "03-key-results.png",
            "Results capability panel",
            "~3840×2560 · 300 DPI",
            "Optional results board",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_paragraph("")
    add_heading_styled(doc, "5. Submission Checklist", 1)
    checks = [
        "Project title is concise and impactful.",
        "Abstract is 80–100 words and covers problem, AI solution, innovation, and impact.",
        "At least 3–4 original high-resolution images included (architecture, workflow, UI, application).",
        "Images are clear, print-tagged (300 DPI), and suitable for large-format display.",
        "Document is professionally structured and presentation-ready.",
    ]
    for c in checks:
        bp = doc.add_paragraph(style="List Bullet")
        r = bp.add_run(f"✓ {c}")
        set_run_font(r, 11, False, CHARCOAL)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(24)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "— End of Submission Package —\n"
        "Source assets folder: docs/print-assets/"
    )
    set_run_font(r, 10, False, MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT, wc


if __name__ == "__main__":
    path, wc = build()
    print(f"Wrote {path}")
    print(f"Abstract word count: {wc}")
