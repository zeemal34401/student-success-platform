#!/usr/bin/env python3
"""Generate a comprehensive technical report (.docx) for the Student Success Platform."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Student_Success_Platform_Technical_Report.docx"


def set_document_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 4):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "Calibri"
        heading.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    run = paragraph.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    run = paragraph.add_run("Right-click and choose Update Field to refresh the table of contents.")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph("")
    return table


def build_report():
    doc = Document()
    set_document_styles(doc)

    # Title page
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Student Success Platform")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Report")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x10, 0x18, 0x28)

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "AI-Based Academic Performance Prediction\n"
        "and Student Success Intelligence Platform\n\n"
        f"Document Version: 1.0\n"
        f"Date: {date.today().strftime('%B %d, %Y')}\n"
        "Classification: Internal Technical Documentation"
    )

    doc.add_page_break()

    # Table of contents
    add_heading(doc, "Table of Contents", 1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "The Student Success Platform (internally referred to as PrognosEd) is an integrated "
        "web application that combines predictive machine learning, role-based dashboards, "
        "intervention workflows, institutional reporting, and a retrieval-augmented (RAG) "
        "advising chatbot. The system is designed to help universities identify at-risk "
        "students early, recommend targeted interventions, and provide faculty and administrators "
        "with actionable intelligence at course, department, and institution-wide levels.",
    )
    add_para(
        doc,
        "The platform implements the full prototype scope described in the AI-Based Academic "
        "Performance Prediction specification: faculty dashboards, risk alerting, a recommendation "
        "engine, institutional analytics, and four independent ML microservices. It extends the "
        "baseline specification with production-oriented features including JWT authentication, "
        "SMTP-based faculty onboarding, password reset, avatar management, global search, and "
        "per-student conversational advising memory.",
    )
    add_bullets(
        doc,
        [
            "Frontend: React 19 with Vite and Tailwind CSS 4",
            "Backend API: Node.js Express with SQLite (better-sqlite3)",
            "ML layer: Four FastAPI Python services (ports 8000–8003)",
            "Advising AI: TF-IDF RAG chat with student-specific conversation history",
            "Deployment: Single-process production mode or Docker Compose",
        ],
    )

    # 2. Introduction
    add_heading(doc, "2. Introduction", 1)
    add_heading(doc, "2.1 Purpose", 2)
    add_para(
        doc,
        "Higher education institutions face increasing pressure to improve retention, graduation "
        "rates, and equitable student outcomes. Manual monitoring of attendance, grades, and "
        "learning-management-system (LMS) engagement does not scale across large cohorts. This "
        "platform automates risk detection, surfaces prioritized student lists, and supports "
        "evidence-based intervention planning for faculty and leadership.",
    )
    add_heading(doc, "2.2 Scope", 2)
    add_para(doc, "This report documents:")
    add_bullets(
        doc,
        [
            "System architecture and technology stack",
            "Functional capabilities by user role",
            "Machine learning models and integration patterns",
            "RAG-based student advising chat",
            "Database schema and data model",
            "REST API surface area",
            "Security, authentication, and authorization",
            "Deployment, configuration, and operational procedures",
            "Testing, limitations, and recommended future work",
        ],
    )
    add_heading(doc, "2.3 Intended Audience", 2)
    add_para(
        doc,
        "Software engineers, ML practitioners, academic administrators, project supervisors, "
        "and stakeholders evaluating the platform for pilot or production deployment.",
    )

    # 3. System Architecture
    add_heading(doc, "3. System Architecture", 1)
    add_heading(doc, "3.1 High-Level Architecture", 2)
    add_para(
        doc,
        "The application follows a three-tier architecture with optional ML microservices:",
    )
    add_numbered(
        doc,
        [
            "Presentation tier — Single-page React application served by Vite in development "
            "and by Express static middleware in production.",
            "Application tier — Express REST API handling authentication, authorization, "
            "business logic, data persistence, email delivery, and ML orchestration.",
            "Analytics tier — Four independent FastAPI services hosting trained scikit-learn "
            "and XGBoost models for academic risk, dropout risk, skill recommendations, and "
            "behavioral clustering.",
            "Knowledge tier — Static JSON advising playbook indexed with TF-IDF for RAG "
            "responses, augmented by per-student chat history stored in SQLite.",
        ],
    )
    add_heading(doc, "3.2 Component Map", 2)
    add_table(
        doc,
        ["Layer", "Technology", "Location"],
        [
            ("Frontend", "React 19, Vite 8, Tailwind CSS 4, Recharts, Lucide icons", "src/"),
            ("API Server", "Express 4, better-sqlite3, bcryptjs, jsonwebtoken", "server/"),
            ("ML Services", "FastAPI, scikit-learn, XGBoost", "ml/model1_pipeline … model4_pipeline/"),
            ("Database", "SQLite (migration-ready for PostgreSQL/MySQL)", "server/data/student_success.db"),
            ("RAG Engine", "Custom TF-IDF retriever (Node.js port)", "server/rag/"),
            ("CI/CD", "GitHub Actions (lint, build, smoke tests)", ".github/workflows/ci.yml"),
        ],
    )
    add_heading(doc, "3.3 Request Flow", 2)
    add_para(
        doc,
        "A typical faculty workflow proceeds as follows: the user authenticates via POST "
        "/api/auth/login with email, password, and selected role; the API returns a JWT stored "
        "client-side. Subsequent requests include Authorization: Bearer <token>. The faculty "
        "dashboard loads aggregated metrics from GET /api/dashboard/faculty. Navigating to "
        "My Students retrieves course-scoped rosters via GET /api/students with optional filters. "
        "Opening a student profile calls GET /api/students/:id, which enriches records with risk "
        "trends and intervention templates. ML predictions are fetched on demand via "
        "GET /api/students/:id/ml-risk. The embedded RAG chat uses GET/POST /api/rag/:studentId.",
    )
    add_heading(doc, "3.4 Resilience and Degradation", 2)
    add_para(
        doc,
        "When ML microservices are unavailable (ECONNREFUSED or timeout), the platform does not "
        "fail entirely. The mlPrediction.service.js module applies heuristic fallbacks derived from "
        "stored student metrics (attendance, GPA, LMS activity, late assignments). The UI surfaces "
        "a clear disclaimer that live model inference is offline. Platform-native skill diagnostics "
        "and rule-based recommendations remain available regardless of ML service status.",
    )

    # 4. Technology Stack
    add_heading(doc, "4. Technology Stack", 1)
    add_table(
        doc,
        ["Category", "Package / Tool", "Version / Notes"],
        [
            ("Runtime", "Node.js", "ES modules (type: module)"),
            ("Frontend framework", "React", "19.2.x"),
            ("Build tool", "Vite", "8.x with @vitejs/plugin-react"),
            ("Styling", "Tailwind CSS", "4.x via @tailwindcss/vite"),
            ("Charts", "Recharts", "3.x"),
            ("HTTP server", "Express", "4.21.x"),
            ("Database driver", "better-sqlite3", "11.x synchronous SQLite"),
            ("Auth", "jsonwebtoken, bcryptjs", "JWT sessions in auth_sessions table"),
            ("Email", "nodemailer", "SMTP faculty invites and password reset"),
            ("Rate limiting", "express-rate-limit", "Login and API throttling"),
            ("Linting", "oxlint", "Fast JS linter"),
            ("ML serving", "FastAPI + uvicorn", "One process per model pipeline"),
        ],
    )

    # 5. User Roles and Features
    add_heading(doc, "5. User Roles and Functional Features", 1)
    add_heading(doc, "5.1 Role Hierarchy", 2)
    add_para(
        doc,
        "The platform implements a hierarchical role model: Director/Dean at the top, "
        "Academic Admin and Administrative Staff at the institutional level, Department Head "
        "(merged into faculty UX), and Faculty at the course level. Each role receives a "
        "tailored navigation menu defined in src/utils/roleScope.js.",
    )
    add_table(
        doc,
        ["Role", "Primary Portal Focus", "Key Screens"],
        [
            (
                "Director / Dean",
                "University-wide performance, department comparison, high-risk drill-down",
                "DirectorDashboard, DirectorRiskDrilldown, DirectorAdminPanel, InstitutionalReports",
            ),
            (
                "Academic Admin",
                "Academic operations, user management, institutional insights",
                "AcademicAdminDashboard, AdminPanel, InstitutionalReports, RiskAlertPanel",
            ),
            (
                "Department Head",
                "Department student performance (faculty-equivalent navigation)",
                "FacultyDashboard, FacultyStudents, FacultyOverview",
            ),
            (
                "Faculty",
                "Own students, attendance, marks, risk alerts, recommendations, RAG chat",
                "FacultyDashboard, FacultyStudents, StudentDetail, RecommendationEngine",
            ),
            (
                "Administrative Staff",
                "Institutional reports and risk alerts (admin-like navigation)",
                "AcademicAdminDashboard, InstitutionalReports, RiskAlertPanel",
            ),
        ],
    )

    add_heading(doc, "5.2 Director / Dean Features", 2)
    add_bullets(
        doc,
        [
            "Executive Dashboard with institution KPIs: total students, critical/high risk counts, retention trends.",
            "Department comparison charts and drill-down into department faculty rosters.",
            "Faculty course and student exploration via nested API endpoints under /api/dashboard/director.",
            "University Reports with export-oriented institutional analytics.",
            "Academic Insights (risk alerts) with cross-department visibility.",
            "Admin Management panel for inviting and managing administrator accounts.",
        ],
    )

    add_heading(doc, "5.3 Academic Admin Features", 2)
    add_bullets(
        doc,
        [
            "Academic Overview dashboard with faculty, course, and enrollment statistics.",
            "Admin Panel for creating faculty accounts, assigning courses, and sending SMTP activation emails.",
            "Email verification (format + MX/mailbox check) before invite dispatch.",
            "Institutional Reports and Academic Insights for cross-cutting analytics.",
            "Global student search with role-scoped results.",
        ],
    )

    add_heading(doc, "5.4 Faculty Features", 2)
    add_para(
        doc,
        "The faculty experience received a dedicated UI overhaul with reusable components in "
        "src/components/faculty/: FacultyHero, FacultyBreadcrumb, FacultyCourseCard, "
        "FacultyAtRiskPanel, and FacultyMetricBar. Styling utilities live in design-system.css.",
    )
    add_bullets(
        doc,
        [
            "Faculty Dashboard — KPI cards (total students, at-risk count, average attendance, LMS engagement), clickable metrics that navigate to filtered student lists, weekly engagement chart (Recharts), and at-risk sidebar panel.",
            "My Students — Course grid with enrollment and risk summaries; drill-down to sortable, filterable rosters; breadcrumb navigation; metric progress bars per student.",
            "Student Detail — Comprehensive profile with risk score, trend sparkline, intervention checklist, ML risk panel, platform recommendations, and embedded RAG advising chat.",
            "Class Reports — Course-scoped reporting for the current term.",
            "Settings — Profile, avatar upload, notification preferences, password change.",
        ],
    )

    add_heading(doc, "5.5 Recommendation Engine", 2)
    add_para(
        doc,
        "The RecommendationEngine screen presents prioritized intervention suggestions derived from "
        "student risk levels and platform heuristics. Faculty can accept or dismiss recommendations; "
        "decisions persist in recommendation_decisions. ML skill recommendations from Model 3 "
        "supplement platform diagnostics when the recommender service is online.",
    )

    add_heading(doc, "5.6 Risk Alerts and Institutional Reports", 2)
    add_para(
        doc,
        "RiskAlertPanel surfaces students above configurable risk thresholds with sorting and "
        "quick navigation to profiles. InstitutionalReports aggregates retention rates, department "
        "risk snapshots, and engagement weekly snapshots for leadership reporting.",
    )

    add_heading(doc, "5.7 Authentication and Onboarding", 2)
    add_bullets(
        doc,
        [
            "Login requires email, password, and explicit role selection to prevent cross-role token misuse.",
            "JWT tokens with configurable expiry (default 24h); sessions tracked in auth_sessions.",
            "Faculty invite flow: admin creates account → SMTP email with activation link → AcceptInvite screen → password set → Active status.",
            "Password reset: forgot-password email with time-limited token (default 30 minutes); dedicated ResetPassword screen and server-side reset-password.html fallback.",
            "Demo accounts available in development only (disabled in production NODE_ENV).",
        ],
    )

    # 6. Machine Learning
    add_heading(doc, "6. Machine Learning Subsystem", 1)
    add_heading(doc, "6.1 Overview", 2)
    add_para(
        doc,
        "Four independent ML pipelines reside under ml/. Each follows a consistent structure: "
        "preprocessing, training scripts, explainability analysis, serialized models in models/, "
        "and a FastAPI api.py entry point. The Node API orchestrates calls via environment-configured URLs.",
    )
    add_table(
        doc,
        ["Model", "Purpose", "Port", "Key Metric", "Primary Dataset"],
        [
            ("Model 1", "Academic early-warning (pass/fail risk)", "8000", "ROC-AUC ~0.83 (4-week window)", "OULAD"),
            ("Model 2", "Dropout risk (Dropout vs Graduate)", "8001", "94.1% accuracy, AUC 0.973", "UCI Dropout"),
            ("Model 3", "Skill intervention recommender", "8002", "Precision@5 100%", "KDD Cup 2010 EDM"),
            ("Model 4", "Behavioral clustering for analytics", "8003", "Silhouette 0.315, k=4", "Multi-source analytics"),
        ],
    )

    add_heading(doc, "6.2 Model 1 — Academic Performance Prediction", 2)
    add_para(
        doc,
        "Predicts early-warning academic risk using engagement signals (LMS clicks, active days) "
        "and assessment performance. Evaluated on a 4-week early-warning variant of OULAD to "
        "reflect realistic intervention lead time. SHAP analysis confirms reliance on behavioral "
        "and academic features rather than demographic proxies.",
    )

    add_heading(doc, "6.3 Model 2 — Dropout Risk Detection", 2)
    add_para(
        doc,
        "Binary classifier distinguishing dropout vs graduate outcomes using enrollment, "
        "financial, and curricular unit features from the UCI Dropout dataset. Achieves high "
        "accuracy with stratified 5-fold cross-validation reported alongside held-out test metrics.",
    )

    add_heading(doc, "6.4 Model 3 — Intervention Recommender", 2)
    add_para(
        doc,
        "Ranks skill-level weak areas from tutoring log mastery gaps. Operates on a pre-computed "
        "mastery table (standard in educational data mining research). Returns top-k skill "
        "recommendations integrated into the student recommendation panel.",
    )

    add_heading(doc, "6.5 Model 4 — Learning Analytics Clustering", 2)
    add_para(
        doc,
        "Unsupervised k-means clustering (k=4) for behavioral segmentation supporting institutional "
        "reporting and cohort analysis. Silhouette score 0.315 indicates moderate cluster separation "
        "appropriate for exploratory analytics.",
    )

    add_heading(doc, "6.6 Integration Layer", 2)
    add_para(
        doc,
        "server/services/mlPrediction.service.js maps platform student records to model-specific "
        "feature vectors. Because live LMS event streams are not ingested in the prototype, some "
        "model features are estimated from stored metrics (attendance → active days, GPA → avg_score, "
        "etc.). Health checks expose service status at GET /api/health/ml.",
    )
    add_table(
        doc,
        ["Environment Variable", "Default", "Service"],
        [
            ("ML_ACADEMIC_API_URL", "http://localhost:8000", "Model 1"),
            ("ML_DROPOUT_API_URL", "http://localhost:8001", "Model 2"),
            ("ML_RECOMMENDER_API_URL", "http://localhost:8002", "Model 3"),
            ("ML_ANALYTICS_API_URL", "http://localhost:8003", "Model 4"),
            ("ML_FETCH_TIMEOUT_MS", "8000", "Request timeout"),
        ],
    )

    add_heading(doc, "6.7 Methodology Notes", 2)
    add_bullets(
        doc,
        [
            "5-fold stratified cross-validation on all classification models.",
            "Class imbalance handled via class-weighting (preferred over SMOTE on smaller datasets).",
            "SHAP explainability for Models 1 and 2.",
            "Early-warning validity: Model 1 uses only first 4 weeks of course data.",
            "Architectural transparency: Model 3 uses pre-computed mastery tables, documented explicitly.",
        ],
    )

    # 7. RAG Chat
    add_heading(doc, "7. RAG-Based Student Advising Chat", 1)
    add_heading(doc, "7.1 Architecture", 2)
    add_para(
        doc,
        "The RAG (Retrieval-Augmented Generation) module provides contextual advising assistance "
        "within each student profile. It combines a static knowledge base "
        "(server/rag/dataset-student-success.json, approximately 30 Q&A pairs covering attendance, "
        "GPA, LMS engagement, outreach templates, and intervention strategies) with per-student "
        "conversation history persisted in rag_chat_messages.",
    )
    add_heading(doc, "7.2 Retrieval Mechanism", 2)
    add_para(
        doc,
        "Unlike the original Python/FAISS reference implementation, production uses a Node.js "
        "TF-IDF retriever (server/rag/retriever.js) for zero-dependency deployment. Documents are "
        "indexed at service startup; queries retrieve top-k relevant playbook chunks.",
    )
    add_heading(doc, "7.3 Answer Generation", 2)
    add_para(
        doc,
        "server/services/ragChat.service.js implements intent detection (outreach, weekly plan, "
        "summary, risk explanation, attendance, assignments, LMS, support, academic) to tailor "
        "responses. Student snapshot data (risk score, trend, metrics) personalizes answers. "
        "Anti-repetition logic rotates focus areas when duplicate questions are detected and skips "
        "recently used playbook chunks. Assistant messages are stored as concise summaries rather "
        "than full response blobs to keep history useful.",
    )
    add_heading(doc, "7.4 Frontend Component", 2)
    add_para(
        doc,
        "StudentRagChat.jsx renders inside StudentDetail.jsx with conversation history above a "
        "prominent compose bar (textarea + Send). Scrolling is confined to the chat container; "
        "opening a profile no longer jumps to the RAG section. Faculty mode enhances the student "
        "header with additional context.",
    )
    add_heading(doc, "7.5 API Endpoints", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Description"],
        [
            ("GET", "/api/rag/:studentId", "Retrieve chat history and student context"),
            ("POST", "/api/rag/:studentId", "Send user message; receive assistant response"),
        ],
    )

    # 8. Database
    add_heading(doc, "8. Database Design", 1)
    add_heading(doc, "8.1 Engine and Migrations", 2)
    add_para(
        doc,
        "SQLite via better-sqlite3 provides synchronous, embedded persistence suitable for "
        "development and pilot deployments. Schema migrations are versioned SQL files in "
        "server/db/schema/ applied by migrate.js. The design maps cleanly to PostgreSQL or MySQL "
        "for production scale-up.",
    )
    add_heading(doc, "8.2 Core Entities", 2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ("departments", "Academic departments with slug identifiers"),
            ("roles", "User role definitions (Director, Admin, Faculty, etc.)"),
            ("terms", "Academic terms with is_current flag"),
            ("users", "Authenticated users with role, department, status (Active/Invited/Disabled)"),
            ("courses", "Course catalog linked to departments"),
            ("faculty_courses", "Many-to-many faculty-course-term assignments"),
            ("students", "Student records linked to courses and departments"),
            ("student_metrics", "Per-term attendance, GPA, LMS activity, late assignments, risk score/level"),
            ("intervention_templates", "Risk-level-specific recommended actions"),
            ("recommendation_decisions", "Faculty accept/dismiss tracking per student-term"),
            ("auth_sessions", "JWT session hashes with expiry and revocation"),
        ],
    )
    add_heading(doc, "8.3 Analytics and Reporting Tables", 2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ("retention_rates", "Historical retention by term"),
            ("department_risk_snapshots", "Risk distribution percentages per department-term"),
            ("engagement_weekly_snapshots", "Weekly attendance/LMS averages by scope (institution/department/faculty)"),
            ("user_notification_preferences", "Per-user alert and digest settings"),
        ],
    )
    add_heading(doc, "8.4 Extended Schema (Migrations 002–007)", 2)
    add_table(
        doc,
        ["Migration", "Feature"],
        [
            ("002_invite_tokens.sql", "Faculty invitation tokens with expiry"),
            ("003_role_hierarchy.sql", "Role hierarchy metadata"),
            ("004_user_avatar.sql", "Avatar file paths for user profiles"),
            ("005_password_reset.sql", "Password reset token storage"),
            ("006_student_courses.sql", "Multi-course enrollment support"),
            ("007_rag_chat.sql", "RAG chat message history per student-user pair"),
        ],
    )
    add_heading(doc, "8.5 Seeding", 2)
    add_para(
        doc,
        "server/db/seed.js populates demo departments, courses, faculty assignments, students with "
        "realistic metrics, intervention templates, retention data, engagement snapshots, and "
        "sample RAG chat threads for demonstration students (STU-1001, STU-1006, STU-1011). "
        "Seeding runs automatically via postinstall (scripts/ensure-ready.js) and can be forced "
        "with npm run db:reset.",
    )

    # 9. API Reference Summary
    add_heading(doc, "9. REST API Reference Summary", 1)
    add_para(doc, "All API routes are prefixed with /api. Authenticated routes require Bearer JWT.")

    add_heading(doc, "9.1 Health and Status", 2)
    add_table(
        doc,
        ["Method", "Path", "Auth", "Description"],
        [
            ("GET", "/health", "No", "API liveness check"),
            ("GET", "/health/ml", "No", "ML microservice connectivity"),
            ("GET", "/health/email", "No", "SMTP configuration status"),
        ],
    )

    add_heading(doc, "9.2 Authentication", 2)
    add_table(
        doc,
        ["Method", "Path", "Description"],
        [
            ("POST", "/auth/login", "Authenticate with email, password, role"),
            ("POST", "/auth/logout", "Revoke current session"),
            ("GET", "/auth/me", "Current user profile"),
            ("GET", "/auth/verify", "Validate token"),
            ("GET", "/auth/demo-accounts", "List demo credentials (dev only)"),
            ("GET", "/auth/invite/:token", "Preview faculty invite"),
            ("POST", "/auth/accept-invite", "Complete faculty activation"),
            ("POST", "/auth/forgot-password", "Request password reset email"),
            ("GET", "/auth/reset/:token", "Preview reset token"),
            ("POST", "/auth/reset-password", "Set new password"),
        ],
    )

    add_heading(doc, "9.3 Students", 2)
    add_table(
        doc,
        ["Method", "Path", "Description"],
        [
            ("GET", "/students", "List students with filters (facultyId, department, riskLevel, search, sort)"),
            ("GET", "/students/search", "Quick search by query string"),
            ("GET", "/students/:id", "Student detail with interventions and risk trend"),
            ("GET", "/students/:id/ml-risk", "ML academic, dropout, and engagement predictions"),
        ],
    )

    add_heading(doc, "9.4 Dashboard", 2)
    add_table(
        doc,
        ["Method", "Path", "Description"],
        [
            ("GET", "/dashboard/faculty", "Faculty dashboard aggregates"),
            ("GET", "/dashboard/department", "Department head dashboard"),
            ("GET", "/dashboard/director", "Director executive dashboard"),
            ("GET", "/dashboard/director/departments", "Department list with metrics"),
            ("GET", "/dashboard/director/departments/:dept/faculty", "Faculty in department"),
            ("GET", "/dashboard/director/faculty/:id/courses", "Faculty course list"),
            ("GET", "/dashboard/director/faculty/:id/students", "Students for faculty course"),
            ("GET", "/dashboard/academic-admin", "Academic admin overview"),
            ("GET", "/dashboard/faculty-overview", "Cross-faculty summary"),
            ("GET", "/dashboard/faculty/:id", "Individual faculty member detail"),
        ],
    )

    add_heading(doc, "9.5 Recommendations, Reports, Admin, Search, Settings", 2)
    add_table(
        doc,
        ["Route Prefix", "Key Operations"],
        [
            ("/recommendations", "GET list; GET /:id/ml-skills; POST /:studentId/decision"),
            ("/reports", "Institutional and departmental report data endpoints"),
            ("/admin", "User CRUD, faculty invites, course assignments, role management"),
            ("/search", "Global scoped search across students, faculty, courses"),
            ("/settings", "Profile update, avatar upload/delete, password change, notification prefs"),
            ("/rag", "GET/POST per-student advising chat"),
        ],
    )

    # 10. Security
    add_heading(doc, "10. Security Considerations", 1)
    add_bullets(
        doc,
        [
            "JWT-based stateless authentication with server-side session tracking for revocation.",
            "bcrypt password hashing with configurable JWT_SECRET (minimum 32 characters in production).",
            "Role-scoped data access enforced in service layer (students.service.js, dashboard.service.js).",
            "Rate limiting on login, password reset, and general API endpoints.",
            "CORS restricted to configured origin (default http://localhost:5173).",
            "Demo accounts automatically disabled when NODE_ENV=production.",
            "SMTP credentials stored server-side only; never exposed to frontend.",
            "Email verification before faculty invites (SMTP or MX mode).",
            "Request body size limit (1MB) and x-powered-by header disabled.",
            "Avatar uploads served from controlled static path with cache headers.",
        ],
    )

    # 11. Frontend Architecture
    add_heading(doc, "11. Frontend Architecture", 1)
    add_para(
        doc,
        "The React SPA uses a single App.jsx orchestrator managing authentication state, role-based "
        "navigation, view routing (without react-router — internal view state machine), loading "
        "skeletons, toast notifications, and global search overlay. API calls centralize in "
        "src/api/client.js with token management and 401 unauthorized handler.",
    )
    add_heading(doc, "11.1 Key Screens", 2)
    add_table(
        doc,
        ["Screen", "File", "Purpose"],
        [
            ("Login", "Login.jsx", "Authentication with role selector and demo panel"),
            ("Faculty Dashboard", "FacultyDashboard.jsx", "KPIs, engagement chart, at-risk panel"),
            ("My Students", "FacultyStudents.jsx", "Course grid and sortable rosters"),
            ("Student Detail", "StudentDetail.jsx", "Profile, ML risk, recommendations, RAG chat"),
            ("Recommendation Engine", "RecommendationEngine.jsx", "Intervention queue with decisions"),
            ("Risk Alerts", "RiskAlertPanel.jsx", "Filtered at-risk student list"),
            ("Director Dashboard", "DirectorDashboard.jsx", "Executive institutional view"),
            ("Admin Panel", "AdminPanel.jsx", "User and faculty management"),
            ("Settings", "Settings.jsx", "Profile, avatar, notifications, password"),
            ("Institutional Reports", "InstitutionalReports.jsx", "Retention and risk analytics"),
        ],
    )
    add_heading(doc, "11.2 Design System", 2)
    add_para(
        doc,
        "src/styles/design-system.css defines CSS custom properties, utility classes, faculty-specific "
        "components (.faculty-hero, .faculty-course-card, .rag-chat-compose), skeleton loaders, and "
        "consistent color palette anchored on primary green (#0B6E4F). Tailwind 4 integrates via "
        "Vite plugin for utility-first styling alongside semantic CSS classes.",
    )

    # 12. Deployment
    add_heading(doc, "12. Deployment and Operations", 1)
    add_heading(doc, "12.1 Development", 2)
    add_para(doc, "Commands:")
    add_bullets(
        doc,
        [
            "npm install — installs dependencies, copies .env, migrates and seeds database",
            "npm run dev (alias npm run app) — starts API on :3001 and Vite on :5173 concurrently",
            "ML services — four uvicorn processes on ports 8000–8003 (see ml/README.md)",
        ],
    )
    add_heading(doc, "12.2 Production", 2)
    add_bullets(
        doc,
        [
            "npm run build — Vite production build to dist/",
            "NODE_ENV=production JWT_SECRET=<secret> npm start — Express serves API + static frontend",
            "docker compose up --build — containerized deployment with environment variables",
        ],
    )
    add_heading(doc, "12.3 Environment Variables", 2)
    add_table(
        doc,
        ["Variable", "Purpose"],
        [
            ("PORT", "API server port (default 3001)"),
            ("JWT_SECRET", "Signing key for tokens (required in production)"),
            ("DATABASE_PATH", "SQLite file location"),
            ("CORS_ORIGIN", "Allowed frontend origin"),
            ("APP_BASE_URL", "Public URL embedded in activation/reset emails"),
            ("SMTP_*", "Email delivery configuration"),
            ("ML_*_API_URL", "ML microservice endpoints"),
            ("VITE_API_URL", "Frontend API base path (default /api)"),
        ],
    )
    add_heading(doc, "12.4 Demo Credentials (Development)", 2)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ("Director / Dean", "director@university.edu", "director123"),
            ("Academic Admin", "admin@university.edu", "admin123"),
            ("Department Head", "head@university.edu", "head123"),
            ("Faculty", "faculty@university.edu", "faculty123"),
            ("Administrative Staff", "staff@university.edu", "staff123"),
        ],
    )

    # 13. Testing
    add_heading(doc, "13. Testing and Quality Assurance", 1)
    add_bullets(
        doc,
        [
            "Smoke tests: npm test runs Node.js native test runner against tests/smoke.test.js",
            "Linting: npm run lint (oxlint)",
            "Build verification: npm run build ensures frontend compiles",
            "CI pipeline: GitHub Actions runs lint, build, and tests on push/PR",
            "Manual QA paths: login per role, dashboard load, student drill-down, RAG chat, admin invite flow",
        ],
    )

    # 14. Spec Compliance
    add_heading(doc, "14. Specification Compliance Matrix", 1)
    add_para(
        doc,
        "The following matrix maps original AI-Based Academic Performance Prediction prototype "
        "requirements to implementation status.",
    )
    add_table(
        doc,
        ["Requirement", "Status", "Implementation Notes"],
        [
            ("Faculty dashboard with at-risk students", "Implemented", "FacultyDashboard + FacultyAtRiskPanel"),
            ("Risk alert system", "Implemented", "RiskAlertPanel with role-scoped filtering"),
            ("Recommendation engine", "Implemented", "Rule-based + ML skill recommendations"),
            ("Institutional reports", "Implemented", "InstitutionalReports with retention/risk snapshots"),
            ("Academic performance ML model", "Implemented", "Model 1 via FastAPI; heuristic fallback offline"),
            ("Dropout risk ML model", "Implemented", "Model 2 via FastAPI; heuristic fallback offline"),
            ("Intervention recommender", "Implemented", "Model 3 + platform skill diagnostic"),
            ("Learning analytics clustering", "Implemented", "Model 4 for behavioral segmentation"),
            ("Multi-source LMS ingestion", "Partial", "Metrics stored in DB; not live event pipeline"),
            ("RAG advising chatbot", "Extended", "TF-IDF RAG with per-student memory (beyond original spec)"),
            ("Role-based access control", "Extended", "Five roles with hierarchical scoping"),
            ("Faculty email onboarding", "Extended", "SMTP invites with mailbox verification"),
        ],
    )

    # 15. Limitations
    add_heading(doc, "15. Known Limitations", 1)
    add_bullets(
        doc,
        [
            "ML feature mapping uses estimated values from summary metrics when raw LMS/assessment events are unavailable.",
            "RAG responses draw from a static advising playbook (~30 entries), not institution-specific policy documents.",
            "SQLite is suitable for pilot/demo; high-concurrency production should migrate to PostgreSQL.",
            "Model 3 requires pre-computed mastery data; arbitrary live skill assessment input is not supported.",
            "Training datasets (~11 GB) are not bundled; retraining requires external dataset access.",
            "RAG does not use a large language model; answers are template-composed from retrieved chunks and student context.",
            "Single-node deployment; horizontal scaling and job queues are not implemented.",
        ],
    )

    # 16. Future Work
    add_heading(doc, "16. Recommended Future Work", 1)
    add_numbered(
        doc,
        [
            "Integrate live LMS and SIS data pipelines (LTI, xAPI, or batch ETL) for ground-truth feature vectors.",
            "Replace or augment TF-IDF RAG with embedding-based retrieval and optional LLM synthesis with citation grounding.",
            "Migrate to PostgreSQL with connection pooling and read replicas for institutional scale.",
            "Add automated model retraining pipelines and model versioning (MLflow or similar).",
            "Implement real-time alerting (email digests, webhook integrations) driven by risk threshold rules.",
            "Expand audit logging and FERPA-compliant data access controls for production compliance.",
            "Add end-to-end and accessibility (WCAG 2.1) test suites.",
            "Containerize ML services with orchestration (Kubernetes) and health-based auto-restart.",
        ],
    )

    # 17. Conclusion
    add_heading(doc, "17. Conclusion", 1)
    add_para(
        doc,
        "The Student Success Platform delivers a cohesive, production-structured prototype that "
        "fulfills the core AI-Based Academic Performance Prediction mandate while adding practical "
        "features for real institutional workflows: secure multi-role access, faculty onboarding, "
        "personalized advising chat, and graceful degradation when ML services are offline. The "
        "modular architecture—React frontend, Express API, independent ML microservices, and "
        "migration-ready SQLite schema—provides a clear path from demonstration to pilot deployment "
        "with targeted investments in data integration, database scaling, and RAG enhancement.",
    )

    # Appendix A
    add_heading(doc, "Appendix A: Project Directory Structure", 1)
    add_para(
        doc,
        "student-success-platform/\n"
        "├── src/                    React frontend (screens, components, api, styles)\n"
        "├── server/                 Express API (routes, services, db, rag, middleware)\n"
        "├── ml/                     Four ML pipelines (model1–model4)\n"
        "├── scripts/                Setup, dev orchestration, report generation\n"
        "├── tests/                  Smoke tests\n"
        "├── dist/                   Production frontend build output\n"
        "├── docker-compose.yml      Container orchestration\n"
        "├── package.json            Node.js dependencies and scripts\n"
        "└── .env.example            Environment variable template",
    )

    # Appendix B
    add_heading(doc, "Appendix B: Dataset Citations", 1)
    add_bullets(
        doc,
        [
            "OULAD: Kuzilek, J., Hlosta, M., Zdrahal, Z. (2017). Open University Learning Analytics Dataset. Scientific Data.",
            "UCI Student Performance: Cortez, P., Silva, A. (2008). UCI Machine Learning Repository.",
            "UCI Dropout: Realinho, V., et al. (2021). DOI: 10.24432/C5MC89.",
            "KDD Cup 2010 EDM Challenge: CMU DataShop.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(f"Report generated: {path}")
